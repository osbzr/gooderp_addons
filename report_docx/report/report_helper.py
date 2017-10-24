# -*- coding: utf-8 -*-
# © 2016 cole
# License AGPL-3.0 or later (http://www.gnu.org/licenses/agpl.html).

from docxtpl import DocxTemplate

import docx

import jinja2

"""
使用一个独立的文件来封装需要支持图片等功能，避免污染report_docx.py
"""


def calc_length(s):
    """
    把字符串，数字类型的参数转化为docx的长度对象，如：
    12 => Pt(12)
    '12' => Pt(12)
    '12pt' => Pt(12)  单位为point
    '12cm' => Cm(12)  单位为厘米
    '12mm' => Mm(12)   单位为毫米
    '12inchs' => Inchs(12)  单位为英寸
    '12emu' => Emu(12)
    '12twips' => Twips(12)
    """
    if not isinstance(s, str):
        # 默认为像素
        return docx.shared.Pt(s)

    if s.endswith('cm'):
        return docx.shared.Cm(float(s[:-2]))
    elif s.endswith('mm'):
        return docx.shared.Mm(float(s[:-2]))
    elif s.endswith('inchs'):
        return docx.shared.Inches(float(s[:-5]))
    elif s.endswith('pt') or s.endswith('px'):
        return docx.shared.Pt(float(s[:-2]))
    elif s.endswith('emu'):
        return docx.shared.Emu(float(s[:-3]))
    elif s.endswith('twips'):
        return docx.shared.Twips(float(s[:-5]))
    else:
        # 默认为像素
        return docx.shared.Pt(float(s))


def calc_alignment(s):
    """
    把字符串转换为对齐的常量
    """
    A = docx.enum.text.WD_ALIGN_PARAGRAPH
    if s == 'center':
        return A.CENTER
    elif s == 'left':
        return A.LEFT
    elif s == 'right':
        return A.RIGHT
    else:
        return A.LEFT


@jinja2.contextfilter
def picture(ctx, data, width=None, height=None, align=None):
    """
    把图片的二进制数据（使用了base64编码）转化为一个docx.Document对象

    data：图片的二进制数据（使用了base64编码）
    width：图片的宽度，可以为：'12cm','12mm','12pt' 等，参考前面的 calc_length()
    height：图片的长度，如果没有设置，根据长度自动缩放
    align：图片的位置，'left'，'center'，'right'
    """

    if not data:
        return None

    # 转化为file-like对象
    # 在python2.7中，bytes==str，可以直接使用
    # 在python3.5中，bytes和str是不同的类型，需要使用base64这个库

    # data使用了base64编码，所以这里需要解码
    data = data.decode('base64')

    import io
    data = io.BytesIO(data)

    tpl = ctx['tpl']
    doc = tpl.new_subdoc()

    if width:
        width = calc_length(width)
    if height:
        height = calc_length(height)

    p = doc.add_paragraph()
    p.alignment = calc_alignment(align)
    p.add_run().add_picture(data, width=width, height=height)
    return doc


def get_env():
    """
    创建一个jinja的enviroment，然后添加一个过滤器 
    """
    jinja_env = jinja2.Environment()
    jinja_env.filters['picture'] = picture
    return jinja_env


def test():
    """
    演示了如何使用，可以直接执行该文件，但是需要使用自己写的docx模版，和图片
    """
    tpl = DocxTemplate("tpls/test_tpl.docx")
    # 读取图片的数据且使用base64编码
    data = open('tpls/python_logo.png', 'rb').read().encode('base64')
    obj = {'logo': data}
    # 需要添加模版对象
    ctx = {'obj': obj, 'tpl': tpl}
    jinja_env = get_env()
    tpl.render(ctx, jinja_env)

    tpl.save('tpls/test.docx')


def main():
    test()


if __name__ == '__main__':
    main()
